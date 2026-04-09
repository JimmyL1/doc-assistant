from pathlib import Path
from langchain_core.documents import Document


def parse_docx(file_path: str | Path) -> list[Document]:
    from docx import Document as DocxDocument
    file_path = Path(file_path)
    doc = DocxDocument(str(file_path))

    # Join non-empty paragraphs into a single text block
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    text = "\n".join(paragraphs)

    return [Document(
        page_content=text,
        metadata={"source": file_path.name, "file_type": "docx", "page": 0},
    )]
